#!/usr/bin/env python3
"""
Convert DOCX files to Markdown format, preserving structure and content.
Handles large documents up to 50+ pages by processing all elements.
"""

import sys
import os
from docx import Document
from docx.document import Document as DocType
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
import re


def get_paragraph_format(paragraph):
    """Extract paragraph formatting information."""
    fmt = paragraph.paragraph_format

    # Determine heading level based on style
    style_name = paragraph.style.name.lower()
    if style_name.startswith('heading'):
        try:
            level = int(re.search(r'\d+', style_name).group())
            return '#' * min(level, 6)  # Limit to H6
        except:
            pass

    return None


def get_text_with_runs(paragraph):
    """Extract text with inline formatting."""
    parts = []
    for run in paragraph.runs:
        text = run.text
        if run.bold and run.italic:
            text = f'***{text}***'
        elif run.bold:
            text = f'**{text}**'
        elif run.italic:
            text = f'*{text}*'
        elif run.underline:
            text = f'<u>{text}</u>'
        parts.append(text)
    return ''.join(parts)


def extract_table_content(table):
    """Extract content from a table."""
    md_table = []
    for i, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            cell_content = []
            for para in cell.paragraphs:
                cell_content.append(get_text_with_runs(para))

            # Process any tables within cells (nested tables)
            for nested_table in cell.tables:
                cell_content.append(extract_table_content(nested_table))

            row_data.append(' '.join(cell_content))

        # Format as markdown table row
        row_str = '| ' + ' | '.join(row_data) + ' |'
        md_table.append(row_str)

        # Add separator after header row
        if i == 0:
            separator = '|' + '|'.join([' --- '] * len(row_data)) + '|'
            md_table.append(separator)

    return '\n'.join(md_table)


def docx_to_markdown(docx_path):
    """Convert a DOCX file to Markdown format."""
    doc = Document(docx_path)

    markdown_lines = []

    # Process each element in the document body
    # Elements are XML nodes (CT_P, CT_Tbl), not Python wrapper objects
    para_tag = qn('w:p')
    table_tag = qn('w:tbl')

    for element in doc.element.body:
        if element.tag == para_tag:
            p = Paragraph(element, doc)

            # Check if it's a heading
            heading_level = get_paragraph_format(p)
            if heading_level:
                markdown_lines.append(f'{heading_level} {get_text_with_runs(p)}')
            else:
                text = get_text_with_runs(p)
                if text.strip():  # Only add non-empty paragraphs
                    markdown_lines.append(text)

        elif element.tag == table_tag:
            t = Table(element, doc.part)
            table_content = extract_table_content(t)
            markdown_lines.append(table_content)

    # Handle elements not directly in the body (headers, footers, textboxes)
    # Process headers
    for section in doc.sections:
        # Header
        header = section.header
        for paragraph in header.paragraphs:
            text = get_text_with_runs(paragraph)
            if text.strip():
                markdown_lines.insert(0, f'> Header: {text}')

        # Footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            text = get_text_with_runs(paragraph)
            if text.strip():
                markdown_lines.append(f'> Footer: {text}')

    return '\n\n'.join(markdown_lines)


def main():
    if len(sys.argv) != 3:
        print("Usage: python convert_docx_to_md.py <input_file.docx> <output_file.md>")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(input_path):
        print(f"Error: Input file '{input_path}' does not exist.")
        sys.exit(1)

    try:
        markdown_content = docx_to_markdown(input_path)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        print(f"Successfully converted '{input_path}' to '{output_path}'")

    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()